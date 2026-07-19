import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Inches

repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
readme_path = os.path.join(repo_root, "README.md")
images_dir = "/Users/cross/Documents/shopflow-1"
out_dir = os.path.join(repo_root, "reports")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "shopflow_report.docx")


def parse_markdown_sections(md_text):
    """Return list of (heading, content_lines) for top-level and sub-headings."""
    sections = []
    current_head = 'Introduction'
    current_lines = []
    for line in md_text.splitlines():
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            # start new section
            sections.append((current_head, current_lines))
            current_head = m.group(2).strip()
            current_lines = []
        else:
            current_lines.append(line)
    sections.append((current_head, current_lines))
    return sections


def add_section_to_doc(doc, heading, lines):
    if heading and heading != 'Introduction':
        doc.add_heading(heading, level=2)
    for line in lines:
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)


def choose_section_for_image(img_name, section_heads):
    name = img_name.lower()
    keywords = [
        'terraform', 'tf', 'eks', 'alb', 'loadbalancer', 'load-balancer', 'load', 'ecr',
        'build', 'push', 'config', 'db', 'database', 'secret', 'deploy', 'helm',
        'storefront', 'catalog', 'orders', 'notifications', 'prometheus', 'fluent',
        'cognito', 'tls', 'ingress', 'alb-insufficient', 'notifications-logged', 'storefront-up'
    ]
    for k in keywords:
        if k in name:
            # pick first section whose heading contains the keyword
            for h in section_heads:
                if k in h.lower():
                    return h
    # fallback: match by any word in heading
    for h in section_heads:
        words = [w for w in re.findall(r"[A-Za-z0-9_-]+", h.lower())]
        for w in words:
            if w in name:
                return h
    return None


doc = Document()
doc.add_heading('ShopFlow Report', level=1)
doc.add_paragraph('This report was generated from README.md and screenshots found in /Users/cross/Documents/shopflow-1.')

sections = []
if os.path.exists(readme_path):
    try:
        with open(readme_path, 'r', encoding='utf-8') as f:
            md = f.read()
        sections = parse_markdown_sections(md)
    except Exception as e:
        doc.add_paragraph(f'Error reading README.md: {e}')
else:
    doc.add_paragraph('README.md not found in repository root.')

# Build map of headings -> list of images
images_map = defaultdict(list)
unspecified = []
if os.path.isdir(images_dir):
    images = sorted([p for p in os.listdir(images_dir) if p.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))])
    section_heads = [h for h, _ in sections]
    for img in images:
        target = choose_section_for_image(img, section_heads)
        if target:
            images_map[target].append(os.path.join(images_dir, img))
        else:
            unspecified.append(os.path.join(images_dir, img))

# Add sections and insert matched images after each section
for heading, lines in sections:
    add_section_to_doc(doc, heading, lines)
    imgs = images_map.get(heading, [])
    for ip in imgs:
        try:
            doc.add_paragraph(os.path.basename(ip))
            doc.add_picture(ip, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f'Could not insert {ip}: {e}')

# Add any unspecified images at the end under Screenshots
if unspecified:
    doc.add_page_break()
    doc.add_heading('Screenshots', level=2)
    for ip in unspecified:
        try:
            doc.add_paragraph(os.path.basename(ip))
            doc.add_picture(ip, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f'Could not insert {ip}: {e}')

try:
    doc.save(out_path)
    print('WROTE', out_path)
except Exception as e:
    print('FAILED to write report:', e)
    print('FAILED to write report:', e)
